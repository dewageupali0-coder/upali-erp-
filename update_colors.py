"""
Update color scheme in Upali Immigration Client Service Agreement v11 -> v12
Old greens -> Navy/Royal Blue scheme with gold accents
"""
from docx import Document
from docx.oxml.ns import qn
import lxml.etree as etree
import copy

SRC = '/root/.claude/uploads/3982b3b7-286e-4bda-8fd8-4c293f73249d/c0d6c7bc-Upali_Immigration_Client_Service_Agreement_v11.docx'
DST = '/home/user/upali-erp-/Upali_Immigration_Client_Service_Agreement_v12.docx'

# Color mapping
COLOR_MAP = {
    # Main header banner (darkest green) -> deep navy
    '0F3A1A': '1B2A6B',
    # Table header cells (mid green) -> royal blue
    '1A5C2A': '1B2A6B',
    # Light green data cells -> very light blue
    'E8F5EA': 'EEF2FF',
    # Regulatory notice background (amber/yellow) -> light amber (keep warm)
    'FFF8E1': 'FFFBEB',
    # Letterhead / service provider cell light green
    'E8F5EA': 'EEF2FF',
}

# For border colors on regulatory notice box: change green borders to amber
BORDER_COLOR_MAP = {
    '1A5C2A': 'F59E0B',
    '0F3A1A': 'F59E0B',
    '2E7D32': 'F59E0B',
    'F9A825': 'F59E0B',
    'FFA000': 'F59E0B',
}

def remap_fill(fill_val):
    if fill_val is None:
        return None
    upper = fill_val.upper()
    mapping = {
        '0F3A1A': '1B2A6B',
        '1A5C2A': '1B2A6B',
        'E8F5EA': 'EEF2FF',
        'FFF8E1': 'FFFBEB',
        'F9F9F9': 'F9F9F9',  # keep as-is
    }
    return mapping.get(upper, fill_val)

def remap_border_color(color_val):
    if color_val is None:
        return None
    upper = color_val.upper()
    # Remap dark green borders to amber
    green_borders = {'1A5C2A', '0F3A1A', '2E7D32', '155724', '1B5E20'}
    amber_borders = {'F9A825', 'FFA000', 'FFB300', 'FFC107', 'FFCA28'}
    if upper in green_borders:
        return 'F59E0B'
    if upper in amber_borders:
        return 'F59E0B'
    return color_val

def update_shd(shd_elem):
    """Update fill and color attributes on a w:shd element."""
    fill = shd_elem.get(qn('w:fill'))
    new_fill = remap_fill(fill)
    if new_fill and new_fill != fill:
        shd_elem.set(qn('w:fill'), new_fill)
        print(f"  shd fill: {fill} -> {new_fill}")
    # Also update w:color on shd if it mirrors fill
    color = shd_elem.get(qn('w:color'))
    if color:
        new_color = remap_fill(color)
        if new_color and new_color != color:
            shd_elem.set(qn('w:color'), new_color)

def update_borders(border_elem):
    """Update border color attributes."""
    for border_tag in ['w:top', 'w:bottom', 'w:left', 'w:right', 'w:insideH', 'w:insideV']:
        b = border_elem.find(qn(border_tag))
        if b is not None:
            color = b.get(qn('w:color'))
            if color:
                new_color = remap_border_color(color)
                if new_color != color:
                    b.set(qn('w:color'), new_color)
                    print(f"  border {border_tag} color: {color} -> {new_color}")

def update_run_color(run_elem):
    """Update run font color if it's a green shade."""
    rPr = run_elem.find(qn('w:rPr'))
    if rPr is None:
        return
    color_elem = rPr.find(qn('w:color'))
    if color_elem is not None:
        val = color_elem.get(qn('w:val'))
        if val:
            upper = val.upper()
            green_shades = {'1A5C2A', '0F3A1A', '2E7D32', '155724', '1B5E20', '388E3C', '4CAF50'}
            if upper in green_shades:
                color_elem.set(qn('w:val'), '1B2A6B')
                print(f"  run color: {val} -> 1B2A6B")
            # Also remap amber/gold run colors to a richer gold
            amber_shades = {'F9A825', 'FFA000', 'FFB300', 'E65100'}
            if upper in amber_shades:
                color_elem.set(qn('w:val'), 'D97706')
                print(f"  run color: {val} -> D97706")

doc = Document(SRC)

print("=== Updating paragraph shading ===")
for i, para in enumerate(doc.paragraphs):
    for shd in para._p.findall('.//' + qn('w:shd')):
        update_shd(shd)
    for run in para._p.findall('.//' + qn('w:r')):
        update_run_color(run)

print("\n=== Updating table cell shading and borders ===")
for t_idx, table in enumerate(doc.tables):
    # Update table-level borders
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is not None:
            print(f"Table {t_idx} tbl borders:")
            update_borders(tblBorders)

    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            tc = cell._tc
            # Cell shading
            for shd in tc.findall('.//' + qn('w:shd')):
                update_shd(shd)
            # Cell borders
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is not None:
                    print(f"Table {t_idx} R{r_idx} C{c_idx} borders:")
                    update_borders(tcBorders)
            # Run colors inside cells
            for run in tc.findall('.//' + qn('w:r')):
                update_run_color(run)

# Also scan all XML for any remaining green hex colors in theme/style parts
print("\n=== Done updating. Saving... ===")
doc.save(DST)
print(f"Saved to {DST}")
