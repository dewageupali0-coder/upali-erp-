from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm
from lxml import etree

doc = Document('Upali_Immigration_Client_Service_Agreement_v13.docx')

# ─── 1. Header banner (Table 0 in body) → bright green ──────────────────────
body_t0 = doc.tables[0]  # "CLIENT SERVICE AGREEMENT" banner
cell = body_t0.rows[0].cells[0]._tc
tcPr = cell.find(qn('w:tcPr'))
shd = tcPr.find(qn('w:shd'))
shd.set(qn('w:fill'), '1A7A3C')   # bright vibrant green

# Update subtitle text color to bright light green
for para in body_t0.rows[0].cells[0].paragraphs:
    for run in para.runs:
        rPr = run._r.find(qn('w:rPr'))
        if rPr is None:
            continue
        color_el = rPr.find(qn('w:color'))
        if color_el is not None:
            val = color_el.get(qn('w:val'), '')
            if val in ('B8E0C0', '1B2A6B', 'FFFFFF'):
                if val == 'FFFFFF':
                    pass  # keep white for title
                else:
                    color_el.set(qn('w:val'), 'A8E6BF')  # light green subtitle

# ─── 2. Info table header cells → bright green ──────────────────────────────
for table in doc.tables[1:]:
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                continue
            shd = tcPr.find(qn('w:shd'))
            if shd is None:
                continue
            fill = shd.get(qn('w:fill'), '')
            if fill in ('1B2A6B', '1A5C2A', '1A5C2E'):
                shd.set(qn('w:fill'), '1A7A3C')  # bright green header cells

# ─── 3. Header logo cell → remove border, clean white background ────────────
header = doc.sections[0].header
ht = header.tables[0]
left_tc = ht.rows[0].cells[0]._tc

# Set logo cell to white (no colored background — removes the "frame" look)
left_tcPr = left_tc.find(qn('w:tcPr'))
shd = left_tcPr.find(qn('w:shd'))
if shd is None:
    shd = OxmlElement('w:shd')
    left_tcPr.append(shd)
shd.set(qn('w:val'),   'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'),  'FFFFFF')   # white — removes colored frame

# Remove any cell borders on logo cell
tcBorders = left_tcPr.find(qn('w:tcBorders'))
if tcBorders is not None:
    left_tcPr.remove(tcBorders)
# Explicitly set all borders to none
tcBorders = OxmlElement('w:tcBorders')
left_tcPr.append(tcBorders)
for side in ('top','left','bottom','right'):
    el = OxmlElement(f'w:{side}')
    el.set(qn('w:val'),   'none')
    el.set(qn('w:sz'),    '0')
    el.set(qn('w:space'), '0')
    el.set(qn('w:color'), 'auto')
    tcBorders.append(el)

# ─── 4. Right header cell → keep light background, update company name to green
right_tc = ht.rows[0].cells[1]._tc
right_tcPr = right_tc.find(qn('w:tcPr'))
rshd = right_tcPr.find(qn('w:shd'))
if rshd is not None:
    rshd.set(qn('w:fill'), 'F2FBF5')   # very light green tint

for para in ht.rows[0].cells[1].paragraphs:
    for run in para.runs:
        rPr = run._r.find(qn('w:rPr'))
        if rPr is None:
            continue
        color_el = rPr.find(qn('w:color'))
        if color_el is not None:
            val = color_el.get(qn('w:val'), '')
            if val in ('1B2A6B', '1A5C2E'):
                color_el.set(qn('w:val'), '1A7A3C')  # bright green company name

# ─── 5. Header table bottom border → bright green ───────────────────────────
tbl = ht._tbl
tblPr = tbl.find(qn('w:tblPr'))
borders_el = tblPr.find(qn('w:tblBorders'))
if borders_el is not None:
    bottom = borders_el.find(qn('w:bottom'))
    if bottom is not None:
        bottom.set(qn('w:color'), '1A7A3C')

doc.save('Upali_Immigration_Client_Service_Agreement_v14.docx')
print("Saved v14")
