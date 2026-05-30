from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

GREEN='1A7A3C'; GREEN_LIGHT='F2FBF5'; GOLD='D4A017'; NAVY='1B2A6B'
WHITE='FFFFFF'; DARK='2D2D2D'; GREY='F7F7F7'; AMBER_BG='FFFBEB'; AMBER_BD='F59E0B'

def shade(cell, color):
    tc=cell._tc; tcPr=tc.find(qn('w:tcPr'))
    if tcPr is None: tcPr=OxmlElement('w:tcPr'); tc.insert(0,tcPr)
    shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),color)

def set_cell_borders(cell, **sides):
    tc=cell._tc; tcPr=tc.find(qn('w:tcPr'))
    if tcPr is None: tcPr=OxmlElement('w:tcPr'); tc.insert(0,tcPr)
    tcB=tcPr.find(qn('w:tcBorders'))
    if tcB is None: tcB=OxmlElement('w:tcBorders'); tcPr.append(tcB)
    for side,(val,sz,color) in sides.items():
        el=tcB.find(qn(f'w:{side}'))
        if el is None: el=OxmlElement(f'w:{side}'); tcB.append(el)
        el.set(qn('w:val'),val); el.set(qn('w:sz'),str(sz))
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),color)

def rPr_color(run, color, bold=False, size=None, italic=False):
    rPr=run._r.find(qn('w:rPr'))
    if rPr is None: rPr=OxmlElement('w:rPr'); run._r.insert(0,rPr)
    c=rPr.find(qn('w:color'))
    if c is None: c=OxmlElement('w:color'); rPr.append(c)
    c.set(qn('w:val'),color)
    if bold:
        b=rPr.find(qn('w:b'))
        if b is None: b=OxmlElement('w:b'); rPr.append(b)
    if size:
        sz=rPr.find(qn('w:sz'))
        if sz is None: sz=OxmlElement('w:sz'); rPr.append(sz)
        sz.set(qn('w:val'),str(size*2))
    if italic:
        i=rPr.find(qn('w:i'))
        if i is None: i=OxmlElement('w:i'); rPr.append(i)

def sp(para, before=0, after=0):
    pPr=para._p.find(qn('w:pPr'))
    if pPr is None: pPr=OxmlElement('w:pPr'); para._p.insert(0,pPr)
    spEl=pPr.find(qn('w:spacing'))
    if spEl is None: spEl=OxmlElement('w:spacing'); pPr.append(spEl)
    spEl.set(qn('w:before'),str(before)); spEl.set(qn('w:after'),str(after))

def set_tbl_borders(table,color=GREEN,sz=8):
    tbl=table._tbl; tblPr=tbl.find(qn('w:tblPr'))
    if tblPr is None: tblPr=OxmlElement('w:tblPr'); tbl.insert(0,tblPr)
    bd=tblPr.find(qn('w:tblBorders'))
    if bd is None: bd=OxmlElement('w:tblBorders'); tblPr.append(bd)
    for s in ('top','left','bottom','right','insideH','insideV'):
        el=bd.find(qn(f'w:{s}'))
        if el is None: el=OxmlElement(f'w:{s}'); bd.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),str(sz))
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),color)

doc = Document()
sec = doc.sections[0]
sec.page_width=Cm(21); sec.page_height=Cm(29.7)
sec.top_margin=Cm(1.5); sec.bottom_margin=Cm(1.8)
sec.left_margin=Cm(2.2); sec.right_margin=Cm(2.0)

# ── Header ───────────────────────────────────────────────────────────────────
hdr = sec.header
ht = hdr.add_table(1, 2, width=Cm(16.8))
ht.style = 'Table Grid'
lc = ht.rows[0].cells[0]; rc = ht.rows[0].cells[1]
ht.columns[0].width = Cm(3.0); ht.columns[1].width = Cm(13.8)
shade(lc, WHITE); shade(rc, GREEN_LIGHT)
set_tbl_borders(ht, color='DDDDDD', sz=4)
lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
lp.add_run().add_picture('/tmp/hdr_logo_bright.png', width=Cm(2.5))
sp(lp, 60, 60)
p1=rc.paragraphs[0]; r1=p1.add_run('UPALI IMMIGRATION SERVICES')
rPr_color(r1,GREEN,bold=True,size=14); sp(p1,80,20)
p2=rc.add_paragraph(); r2=p2.add_run('Immigration Solutions  •  Logistics Connecting Worlds')
rPr_color(r2,GOLD,italic=True,size=9); sp(p2,0,14)
p3=rc.add_paragraph(); r3=p3.add_run('Licence No. 2646027.01   |   TRN: 105425790000001')
rPr_color(r3,DARK,bold=True,size=8); sp(p3,0,8)
p4=rc.add_paragraph(); r4=p4.add_run('Shams Free Zone, Sharjah Media City, Sharjah, UAE')
rPr_color(r4,'555555',size=8); sp(p4,0,8)
p5=rc.add_paragraph(); r5=p5.add_run('+971 54 204 0298  •  info@myupali.lk  •  www.myupali.lk')
rPr_color(r5,'555555',size=8); sp(p5,0,60)
hdr.add_paragraph()

# ── Banner ───────────────────────────────────────────────────────────────────
bt = doc.add_table(1,1); bt.style='Table Grid'; set_tbl_borders(bt,color=GREEN,sz=0)
bc = bt.rows[0].cells[0]; shade(bc, GREEN)
bp = bc.paragraphs[0]; bp.alignment=WD_ALIGN_PARAGRAPH.CENTER
br1 = bp.add_run('WELCOME LETTER'); rPr_color(br1,WHITE,bold=True,size=16); sp(bp,120,40)
bp2 = bc.add_paragraph(); bp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
br2 = bp2.add_run('Confirmation of Engagement — UAE Immigration Consultancy Services')
rPr_color(br2,'A8E6BF',size=9,italic=True); sp(bp2,0,100)
doc.add_paragraph()

# ── Date / Ref block ─────────────────────────────────────────────────────────
rt = doc.add_table(1,2); rt.style='Table Grid'; set_tbl_borders(rt,color='CCCCCC',sz=4)
rl=rt.rows[0].cells[0]; rr=rt.rows[0].cells[1]
shade(rl,GREEN_LIGHT); shade(rr,GREY)
for cell, lines in [
    (rl, [('Date:', '${System.Today}'),('Ref No.:','UIC-LTR-${System.Today.Year}-${Contacts.Folder_ID}'),
          ('Agreement Ref:','${Contacts.Agreement_Reference}')]),
    (rr, [('To:','${Contacts.Salutation} ${Contacts.First_Name} ${Contacts.Last_Name}'),
          ('Passport No.:','${Contacts.Passport_Number}'),
          ('Country:','${Contacts.Country_of_Residence}'),
          ('Email:','${Contacts.Email}'),('Phone:','${Contacts.Phone}')]),
]:
    first=True
    for lbl, val in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first=False
        rl1=p.add_run(f'{lbl}  '); rPr_color(rl1,GREEN,bold=True,size=9)
        rl2=p.add_run(val); rPr_color(rl2,DARK,size=9)
        sp(p, 60 if lbl=='Date:' or lbl=='To:' else 0, 30)

doc.add_paragraph()

# ── Body ─────────────────────────────────────────────────────────────────────
def body_para(doc, text, bold=False, color=DARK, size=10, before=0, after=120):
    p = doc.add_paragraph()
    r = p.add_run(text)
    rPr_color(r, color, bold=bold, size=size)
    sp(p, before, after)
    return p

body_para(doc,'Dear ${Contacts.Salutation} ${Contacts.First_Name},', bold=False, before=0, after=160)

body_para(doc,
    'On behalf of UPALI IMMIGRATION SERVICES, it is my great pleasure to congratulate you on taking this important step towards your new life and career in ${Contacts.Destination_Country}. We are honoured that you have chosen us to guide you through your work visa journey, and we are confident that — with your commitment and our experience — your application will succeed.',
    after=120)

body_para(doc,
    'This letter formally confirms that we have received your initial details and that your file has been opened under Agreement Reference No. ${Contacts.Agreement_Reference}.',
    after=120)

body_para(doc,
    'Your end-to-end ${Contacts.Destination_Country} ${Contacts.Visa_Type} package has been carefully structured in line with UAE Federal Tax Authority (FTA) requirements at a total of AED ${Contacts.Package_Price_AED}. The full breakdown — consultancy fee, visa application fee, embassy appointment, and international air ticket — is detailed in the attached Service Agreement and Tax Invoice.',
    after=160)

# Steps heading
sh = doc.add_paragraph()
shr = sh.add_run('What Happens Next')
rPr_color(shr, GREEN, bold=True, size=11)
sp(sh, 80, 60)

steps = [
    ('1.', 'Review', 'the Service Agreement (${Contacts.Agreement_Reference}) and Tax Invoice (${Contacts.Invoice_Number}) carefully.'),
    ('2.', 'Sign and return', 'the Service Agreement on or before ${System.AddDays.7}. Late signing may delay submission to ${Contacts.Destination_Country} Immigration.'),
    ('3.', 'Submit VAT zero-rating documents:', '(a) passport copy showing no UAE residency visa; (b) proof of ${Contacts.Country_of_Residence} residential address (not older than 3 months); (c) signed non-presence declaration; (d) signed Service Agreement. Required under Art. 31, UAE VAT Executive Regulations.'),
    ('4.', 'Make the initial payment', 'as set out in clause 5.7 of the Service Agreement (AED ${Contacts.Initial_Payment_AED}).'),
    ('5.', 'Case officer contact', '— once signing and initial payment are confirmed, our case officer will contact you to begin document collection for the ${Contacts.Destination_Country} ${Contacts.Visa_Type} application.'),
]
for num, bold_text, rest in steps:
    p = doc.add_paragraph()
    rn = p.add_run(f'{num}  '); rPr_color(rn, GREEN, bold=True, size=10)
    rb = p.add_run(f'{bold_text}  '); rPr_color(rb, DARK, bold=True, size=10)
    rr = p.add_run(rest); rPr_color(rr, DARK, size=10)
    sp(p, 0, 80)

doc.add_paragraph()
body_para(doc,
    'Please remember that, although we will do everything in our power to support a successful outcome, the final decision on any visa application rests exclusively with the ${Contacts.Destination_Country} Immigration authorities, in accordance with UAE Federal Decree-Law No. 29 of 2021 on Entry and Residence of Foreigners. We will keep you fully informed at each step.',
    after=120)

body_para(doc,
    'Should you have any questions about the agreement, the invoice, or the supporting documents required, please contact us at info@myupali.lk or +971 54 204 0298 (also on WhatsApp). Our team is ready to support you.',
    after=120)

body_para(doc,
    'Once again, congratulations on this exciting milestone. We look forward to walking this journey with you.',
    after=160)

# ── Action Required Box ───────────────────────────────────────────────────────
at = doc.add_table(1,1); at.style='Table Grid'
shade(at.rows[0].cells[0], AMBER_BG)
set_cell_borders(at.rows[0].cells[0],
    top=('single',16,GOLD),left=('single',16,GOLD),
    bottom=('single',16,GOLD),right=('single',16,GOLD))
ap = at.rows[0].cells[0].paragraphs[0]
ar1 = ap.add_run('⚠  ACTION REQUIRED WITHIN 7 DAYS   ')
rPr_color(ar1, GOLD, bold=True, size=10)
ar2 = ap.add_run('Please sign and return the Service Agreement (${Contacts.Agreement_Reference}) to info@myupali.lk within 7 days of this letter to keep your application on schedule. Failure to sign within this period may result in cancellation of your file per clause 4.3 of the Agreement.')
rPr_color(ar2, DARK, size=9)
sp(ap, 100, 100)
doc.add_paragraph()

# ── Enclosures ────────────────────────────────────────────────────────────────
ep = doc.add_paragraph()
er = ep.add_run('Enclosures:')
rPr_color(er, GREEN, bold=True, size=10)
sp(ep, 80, 40)
for enc in ['Client Service Agreement (${Contacts.Agreement_Reference})',
            'Tax Invoice (${Contacts.Invoice_Number})']:
    p = doc.add_paragraph()
    r = p.add_run(f'  •  {enc}')
    rPr_color(r, DARK, size=10)
    sp(p, 0, 30)

doc.add_paragraph()

# ── Signature block ───────────────────────────────────────────────────────────
st = doc.add_table(1,2); st.style='Table Grid'; set_tbl_borders(st,color='CCCCCC',sz=4)
sl=st.rows[0].cells[0]; sr=st.rows[0].cells[1]
shade(sl,GREEN_LIGHT); shade(sr,GREY)

p1=sl.paragraphs[0]; r=p1.add_run('Warm regards,')
rPr_color(r, DARK, size=10); sp(p1, 80, 60)
# stamp
stp=sl.add_paragraph(); stp.alignment=WD_ALIGN_PARAGRAPH.LEFT
stp.add_run().add_picture('/tmp/digital_stamp.png', width=Inches(1.3))
sp(stp,20,20)
for line in ['Upali Dewage','Managing Director','UPALI IMMIGRATION SERVICES','+971 54 204 0298  |  info@myupali.lk']:
    p=sl.add_paragraph(); r=p.add_run(line)
    rPr_color(r, GREEN if 'UPALI' in line else DARK, bold='UPALI' in line, size=9)
    sp(p,0,20)

p_r=sr.paragraphs[0]; rr=p_r.add_run('FOR OFFICIAL USE ONLY')
rPr_color(rr, NAVY, bold=True, size=9); sp(p_r,80,60)
for line in ['File opened by: _______________','Date: _______________','Case Officer: _______________','Status: _______________']:
    p=sr.add_paragraph(); r=p.add_run(line)
    rPr_color(r,'555555',size=9); sp(p,0,40)

# ── Navy footer ───────────────────────────────────────────────────────────────
doc.add_paragraph()
ftt=doc.add_table(1,1); ftt.style='Table Grid'
shade(ftt.rows[0].cells[0],'1B2A6B')
fp=ftt.rows[0].cells[0].paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
fr=fp.add_run('UPALI IMMIGRATION SERVICES — Private Consultancy Licensed by Shams Free Zone, UAE — NOT an affiliate of ICP/ICA, GDRFA, or any government body — All visa decisions rest with the relevant immigration authority — No outcome guaranteed')
rPr_color(fr,'A8C8FF',size=8); sp(fp,80,80)

doc.save('/home/user/upali-erp-/UIS_Welcome_Letter_Template_v1.docx')
print("Welcome letter saved")
