from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import zipfile, shutil, io

# ── Colors ──────────────────────────────────────────────────────────────────
GREEN      = '1A7A3C'
GREEN_LIGHT= 'F2FBF5'
GOLD       = 'D4A017'
NAVY       = '1B2A6B'
WHITE      = 'FFFFFF'
DARK       = '2D2D2D'
GREY       = 'F7F7F7'
AMBER_BG   = 'FFFBEB'
AMBER_BD   = 'F59E0B'

def shade(cell, color):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr'); tc.insert(0, tcPr)
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),color)

def set_cell_borders(cell, **sides):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr'); tc.insert(0, tcPr)
    tcB = tcPr.find(qn('w:tcBorders'))
    if tcB is None:
        tcB = OxmlElement('w:tcBorders'); tcPr.append(tcB)
    for side, (val, sz, color) in sides.items():
        el = tcB.find(qn(f'w:{side}'))
        if el is None:
            el = OxmlElement(f'w:{side}'); tcB.append(el)
        el.set(qn('w:val'), val); el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'),'0'); el.set(qn('w:color'), color)

def rPr_color(run, color, bold=False, size=None, italic=False):
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr'); run._r.insert(0, rPr)
    c = rPr.find(qn('w:color'))
    if c is None:
        c = OxmlElement('w:color'); rPr.append(c)
    c.set(qn('w:val'), color)
    if bold:
        b = rPr.find(qn('w:b'))
        if b is None: b = OxmlElement('w:b'); rPr.append(b)
    if size:
        sz = rPr.find(qn('w:sz'))
        if sz is None: sz = OxmlElement('w:sz'); rPr.append(sz)
        sz.set(qn('w:val'), str(size*2))
    if italic:
        i = rPr.find(qn('w:i'))
        if i is None: i = OxmlElement('w:i'); rPr.append(i)

def para_spacing(para, before=0, after=0):
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr'); para._p.insert(0,pPr)
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing'); pPr.append(sp)
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'),  str(after))

def set_tbl_borders(table, color=GREEN, sz=8):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    bd = tblPr.find(qn('w:tblBorders'))
    if bd is None:
        bd = OxmlElement('w:tblBorders'); tblPr.append(bd)
    for s in ('top','left','bottom','right','insideH','insideV'):
        el = bd.find(qn(f'w:{s}'))
        if el is None: el = OxmlElement(f'w:{s}'); bd.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),str(sz))
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),color)

def add_banner(doc, text, subtitle=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    c = tbl.rows[0].cells[0]
    shade(c, GREEN)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    rPr_color(r, WHITE, bold=True, size=16)
    para_spacing(p, 120, 40)
    if subtitle:
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        rPr_color(r2, 'A8E6BF', size=9, italic=True)
        para_spacing(p2, 0, 100)
    # Remove table outer border
    set_tbl_borders(tbl, color=GREEN, sz=0)
    return tbl

# ════════════════════════════════════════════════════════════════════════════
# BUILD INVOICE
# ════════════════════════════════════════════════════════════════════════════
doc = Document()
# Page setup
sec = doc.sections[0]
sec.page_width  = Cm(21); sec.page_height = Cm(29.7)
sec.top_margin  = Cm(1.5); sec.bottom_margin = Cm(1.8)
sec.left_margin = Cm(2.2); sec.right_margin  = Cm(2.0)

# ── Letterhead header (reuse logo from v16 docx) ────────────────────────────
hdr = sec.header
hdr_tbl = hdr.add_table(1, 2, width=Cm(16.8))
hdr_tbl.style = 'Table Grid'
lc = hdr_tbl.rows[0].cells[0]
rc = hdr_tbl.rows[0].cells[1]
hdr_tbl.columns[0].width = Cm(3.0)
hdr_tbl.columns[1].width = Cm(13.8)
shade(lc, WHITE)
shade(rc, GREEN_LIGHT)
set_tbl_borders(hdr_tbl, color='DDDDDD', sz=4)

# Add logo to left cell
lp = lc.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
lrun = lp.add_run()
# We'll paste logo from the v16 docx via zip
lrun.add_picture('/tmp/hdr_logo_bright.png', width=Cm(2.5))
para_spacing(lp, 60, 60)

# Right cell text
rc_p1 = rc.paragraphs[0]
r1 = rc_p1.add_run('UPALI IMMIGRATION SERVICES')
rPr_color(r1, GREEN, bold=True, size=14)
para_spacing(rc_p1, 80, 20)

rc_p2 = rc.add_paragraph()
r2 = rc_p2.add_run('Immigration Solutions  •  Logistics Connecting Worlds')
rPr_color(r2, GOLD, italic=True, size=9)
para_spacing(rc_p2, 0, 16)

rc_p3 = rc.add_paragraph()
r3 = rc_p3.add_run('Licence No. 2646027.01   |   TRN: 105425790000001')
rPr_color(r3, DARK, bold=True, size=8)
para_spacing(rc_p3, 0, 10)

rc_p4 = rc.add_paragraph()
r4 = rc_p4.add_run('Shams Free Zone, Sharjah Media City, Sharjah, UAE')
rPr_color(r4, '555555', size=8)
para_spacing(rc_p4, 0, 8)

rc_p5 = rc.add_paragraph()
r5 = rc_p5.add_run('+971 54 204 0298  •  info@myupali.lk  •  www.myupali.lk')
rPr_color(r5, '555555', size=8)
para_spacing(rc_p5, 0, 60)

# Gold separator line in header
hdr.add_paragraph()

# ── Banner ───────────────────────────────────────────────────────────────────
add_banner(doc, 'TAX INVOICE',
           'Issued pursuant to UAE VAT Decree-Law No. 8 of 2017')
doc.add_paragraph()

# ── Invoice meta + Bill-To two-column table ──────────────────────────────────
meta = doc.add_table(rows=1, cols=2)
meta.style = 'Table Grid'
set_tbl_borders(meta, color='CCCCCC', sz=4)
mc_l = meta.rows[0].cells[0]
mc_r = meta.rows[0].cells[1]
shade(mc_l, GREEN_LIGHT)
shade(mc_r, GREY)

# Left: Bill To
def add_label_value(cell, label, value, first=False):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    r_lbl = p.add_run(f'{label}: ')
    rPr_color(r_lbl, GREEN, bold=True, size=9)
    r_val = p.add_run(value)
    rPr_color(r_val, DARK, size=9)
    para_spacing(p, 0 if not first else 60, 30)

add_label_value(mc_l, 'BILL TO', '', first=True)
p0 = mc_l.paragraphs[0]
rPr_color(p0.runs[0], GREEN, bold=True, size=11)
for lbl, val in [
    ('Client', '${Contacts.First_Name} ${Contacts.Last_Name}'),
    ('Passport No.', '${Contacts.Passport_Number}'),
    ('Country of Residence', '${Contacts.Country_of_Residence}'),
    ('Email', '${Contacts.Email}'),
    ('Phone', '${Contacts.Phone}'),
]:
    add_label_value(mc_l, lbl, val)

# Right: Invoice details
add_label_value(mc_r, 'Invoice No.', 'UIC-INV-${System.NextInvoiceNo}', first=True)
for lbl, val in [
    ('Date of Issue', '${System.Today}'),
    ('Agreement Ref.', '${Contacts.Agreement_Reference}'),
    ('VAT Treatment', 'Zero-Rated (Art. 31)' ),
    ('Client Type', '${Contacts.Client_Type}'),
    ('Due Date', '${System.Today}'),
]:
    add_label_value(mc_r, lbl, val)

doc.add_paragraph()

# ── VAT Notice box ───────────────────────────────────────────────────────────
vat_tbl = doc.add_table(rows=1, cols=1)
vat_tbl.style = 'Table Grid'
vc = vat_tbl.rows[0].cells[0]
shade(vc, AMBER_BG)
set_cell_borders(vc,
    top   =('single', 12, GOLD),
    left  =('single', 12, GOLD),
    bottom=('single', 12, GOLD),
    right =('single', 12, GOLD),
)
vp = vc.paragraphs[0]
vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
vr1 = vp.add_run('⚠  VAT TREATMENT NOTICE   ')
rPr_color(vr1, GOLD, bold=True, size=9)
vr2 = vp.add_run('Consultancy Fee is Zero-Rated (0% VAT) under Article 31, UAE VAT Executive Regulations — export of services to client residing outside UAE. Zero-rating applies only where: (i) client resides outside UAE; (ii) no UAE-based services delivered; (iii) zero-rating documents submitted. Government fees are out-of-scope disbursements. Air ticket zero-rated under Article 45(3).')
rPr_color(vr2, DARK, size=9)
para_spacing(vp, 80, 80)
doc.add_paragraph()

# ── Line Items table ─────────────────────────────────────────────────────────
headers = ['#', 'Description', 'Amount (AED)', 'VAT Rate', 'VAT (AED)', 'Line Total (AED)']
items_tbl = doc.add_table(rows=1, cols=6)
items_tbl.style = 'Table Grid'
set_tbl_borders(items_tbl, color=GREEN, sz=6)

# Header row
hrow = items_tbl.rows[0]
for i, h in enumerate(headers):
    c = hrow.cells[i]
    shade(c, GREEN)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    rPr_color(r, WHITE, bold=True, size=9)
    para_spacing(p, 80, 80)

# Data rows
line_items = [
    ('1', 'Consultancy Fee — ${Contacts.Destination_Country} ${Contacts.Visa_Type}',
     '${Contacts.Consultancy_Fee_AED}', 'Zero-Rated (0%)', '0.00', '${Contacts.Consultancy_Fee_AED}'),
    ('2', 'Visa Application Fee (Govt. Disbursement — pass-through, no mark-up)',
     '${Contacts.Visa_Fee_AED}', 'Out of Scope', '—', '${Contacts.Visa_Fee_AED}'),
    ('3', 'Biometric / Embassy Appointment Fee (Govt. Disbursement)',
     '${Contacts.Biometric_Fee_AED}', 'Out of Scope', '—', '${Contacts.Biometric_Fee_AED}'),
    ('4', 'International Air Ticket (UAE → ${Contacts.Destination_Country}) — Art. 45(3)',
     '${Contacts.Air_Ticket_AED}', 'Zero-Rated (0%)', '0.00', '${Contacts.Air_Ticket_AED}'),
]
for idx, row_data in enumerate(line_items):
    row = items_tbl.add_row()
    bg = GREY if idx % 2 == 0 else WHITE
    for j, val in enumerate(row_data):
        c = row.cells[j]
        shade(c, bg)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(val)
        rPr_color(r, DARK, size=9)
        para_spacing(p, 60, 60)

# Totals rows
totals = [
    ('', '', 'Sub-Total (AED)', '', '', '${Invoice.SubTotal}'),
    ('', '', 'VAT Total (AED)', '', '', '0.00'),
    ('', '', 'TOTAL DUE (AED)', '', '', '${Invoice.Total}'),
]
for i, row_data in enumerate(totals):
    row = items_tbl.add_row()
    is_total = (i == 2)
    for j, val in enumerate(row_data):
        c = row.cells[j]
        shade(c, GREEN if is_total else GREEN_LIGHT)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j >= 2 else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val)
        rPr_color(r, WHITE if is_total else GREEN, bold=is_total, size=10 if is_total else 9)
        para_spacing(p, 80, 80)

doc.add_paragraph()

# ── Payment details ──────────────────────────────────────────────────────────
pay_tbl = doc.add_table(rows=1, cols=1)
pay_tbl.style = 'Table Grid'
shade(pay_tbl.rows[0].cells[0], GREEN_LIGHT)
set_cell_borders(pay_tbl.rows[0].cells[0],
    top=('single',12,GREEN), left=('single',4,'CCCCCC'),
    bottom=('single',4,'CCCCCC'), right=('single',4,'CCCCCC'))
pp = pay_tbl.rows[0].cells[0].paragraphs[0]
pr1 = pp.add_run('PAYMENT SCHEDULE   ')
rPr_color(pr1, GREEN, bold=True, size=10)
pr2 = pp.add_run('(a) AED ${Contacts.Initial_Payment_AED} on signing;  (b) AED ${Contacts.Second_Payment_AED} upon visa application submission;  (c) AED ${Contacts.Final_Payment_AED} upon visa grant.  Bank: ${Company.BankDetails}')
rPr_color(pr2, DARK, size=9)
para_spacing(pp, 80, 80)

doc.add_paragraph()

# ── Signature blocks ─────────────────────────────────────────────────────────
sig_tbl = doc.add_table(rows=1, cols=2)
sig_tbl.style = 'Table Grid'
set_tbl_borders(sig_tbl, color='CCCCCC', sz=4)
sl = sig_tbl.rows[0].cells[0]
sr = sig_tbl.rows[0].cells[1]
shade(sl, GREEN_LIGHT)
shade(sr, GREY)

for cell, title, name, role in [
    (sl, 'FOR UPALI IMMIGRATION SERVICES', 'Upali Dewage', 'Managing Director'),
    (sr, 'RECEIVED BY CLIENT', '${Contacts.First_Name} ${Contacts.Last_Name}', 'Client Signature & Date'),
]:
    p1 = cell.paragraphs[0]
    r = p1.add_run(title)
    rPr_color(r, GREEN, bold=True, size=9)
    para_spacing(p1, 80, 60)
    # Add stamp to service provider
    if 'UPALI' in title:
        sp = cell.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.add_run().add_picture('/tmp/digital_stamp.png', width=Inches(1.4))
        para_spacing(sp, 20, 20)
    p2 = cell.add_paragraph()
    p2.add_run('Signature: ___________________________')
    para_spacing(p2, 80, 20)
    p3 = cell.add_paragraph()
    p3.add_run(f'Name:  {name}')
    para_spacing(p3, 0, 20)
    p4 = cell.add_paragraph()
    p4.add_run(f'Title:  {role}')
    para_spacing(p4, 0, 20)
    p5 = cell.add_paragraph()
    p5.add_run('Date:  ___________________________')
    para_spacing(p5, 0, 80)

doc.add_paragraph()

# ── Legal Footer ─────────────────────────────────────────────────────────────
ft_tbl = doc.add_table(rows=1, cols=1)
ft_tbl.style = 'Table Grid'
shade(ft_tbl.rows[0].cells[0], '1B2A6B')
fp = ft_tbl.rows[0].cells[0].paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run('Invoice issued pursuant to UAE Federal Decree-Law No. 8 of 2017 on VAT. Consultancy fee zero-rated under Article 31 (export of services). Government fees out-of-scope disbursements. Air ticket zero-rated under Article 45(3). Records retained 5 years per Article 78. TRN: 105425790000001')
rPr_color(fr, 'A8C8FF', size=8)
para_spacing(fp, 80, 80)

doc.save('/home/user/upali-erp-/UIS_Tax_Invoice_Template_v1.docx')
print("Invoice saved")
