from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from lxml import etree
import copy

doc = Document('Upali_Immigration_Client_Service_Agreement_v12.docx')

# ─── 1. FIX PAGE MARGINS (reduce top margin so header isn't cramped) ───────────
section = doc.sections[0]
section.top_margin    = Cm(1.5)   # was 3.5cm
section.bottom_margin = Cm(1.8)
section.left_margin   = Cm(2.2)
section.right_margin  = Cm(2.0)

# ─── 2. RESTYLE HEADER TABLE ────────────────────────────────────────────────────
header = section.header
ht = header.tables[0]
tbl = ht._tbl

# ── a) Remove the outer table border (was already none, keep it) ──
# ── b) Give the logo cell a clean navy blue background matching the scheme ──
left_tc  = ht.rows[0].cells[0]._tc
right_tc = ht.rows[0].cells[1]._tc

# Logo cell: navy blue, no border
def set_cell_fill(tc, color_hex):
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_hex)

set_cell_fill(left_tc,  '1B2A6B')   # navy – matches header banner
set_cell_fill(right_tc, 'F0F4FF')   # very light blue-white

# ── c) Update right-cell text colors to navy theme ──
for para in ht.rows[0].cells[1].paragraphs:
    for run in para.runs:
        rPr = run._r.find(qn('w:rPr'))
        if rPr is None:
            continue
        color_el = rPr.find(qn('w:color'))
        if color_el is not None:
            old = color_el.get(qn('w:val'), '')
            if old in ('1A5C2E',):            # company name: was green
                color_el.set(qn('w:val'), '1B2A6B')
            elif old in ('7A6010',):           # tagline: was gold/olive
                color_el.set(qn('w:val'), 'D4A017')   # keep gold accent

# ── d) Add a thin bottom border line under the header table ──
tblPr = tbl.find(qn('w:tblPr'))
borders_el = tblPr.find(qn('w:tblBorders'))
if borders_el is None:
    borders_el = OxmlElement('w:tblBorders')
    tblPr.append(borders_el)

def set_border(borders, side, val='single', sz='12', color='1B2A6B'):
    el = borders.find(qn(f'w:{side}'))
    if el is None:
        el = OxmlElement(f'w:{side}')
        borders.append(el)
    el.set(qn('w:val'),   val)
    el.set(qn('w:sz'),    sz)
    el.set(qn('w:space'), '0')
    el.set(qn('w:color'), color)

# Only add bottom border; keep everything else none
for side in ('top','left','right','insideH','insideV'):
    set_border(borders_el, side, val='none', sz='0', color='auto')
set_border(borders_el, 'bottom', val='single', sz='16', color='D4A017')  # gold underline

# ─── 3. FIX HEADER TABLE WIDTH TO FULL PAGE ────────────────────────────────────
tblW = tblPr.find(qn('w:tblW'))
if tblW is None:
    tblW = OxmlElement('w:tblW')
    tblPr.insert(0, tblW)
tblW.set(qn('w:w'),    '9638')   # full page width in twips (approx A4 minus margins)
tblW.set(qn('w:type'), 'dxa')

# ─── 4. REMOVE LOGO CELL BORDER / ROUND THE IMAGE VISUALLY ────────────────────
# Remove any cell-level borders on the logo cell
left_tcPr = left_tc.find(qn('w:tcPr'))
tcBorders = left_tcPr.find(qn('w:tcBorders'))
if tcBorders is not None:
    left_tcPr.remove(tcBorders)

# ─── 5. ADD SPACING PARAGRAPH AFTER HEADER IN BODY ────────────────────────────
# Ensure first body paragraph has small space before first table
body = doc.element.body
first_para = body.findall(qn('w:p'))[0]
pPr = first_para.find(qn('w:pPr'))
if pPr is None:
    pPr = OxmlElement('w:pPr')
    first_para.insert(0, pPr)
spacing = pPr.find(qn('w:spacing'))
if spacing is None:
    spacing = OxmlElement('w:spacing')
    pPr.append(spacing)
spacing.set(qn('w:before'), '80')
spacing.set(qn('w:after'),  '80')

# ─── 6. SAVE ──────────────────────────────────────────────────────────────────
doc.save('Upali_Immigration_Client_Service_Agreement_v13.docx')
print("Done — saved v13")
